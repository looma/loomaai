from docx import Document
import npttf2utf
from base.fontmapper import FontMapper

def convert(element:str):
    mapper = FontMapper("./files/map.json")
    mappedText = mapper.map_to_unicode(element, from_font="Preeti", unescape_html_input=False, escape_html_output=False)

    return mappedText

def replace_text_in_docx():
    # Open the .docx file
    doc = Document("./files/pdf.docx")
    
    # Iterate through paragraphs in the document
    for para in doc.paragraphs:
        para.text = para.text.replace(para.text, convert(para.text))
    
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
            """

    # Save the changes to a new file
    new_doc_path = 'updated_' + "pdf.docx"
    doc.save(new_doc_path)
    print(f"Updated document saved as {new_doc_path}")

def replace_text_in_place(doc_path):
    # Open the .docx file
    doc = Document(doc_path)
    
    # Function to replace text in a run
    def replace_in_run(para):
        para.text = para.text.replace(para.text, convert(para.text))
    
    # Iterate through paragraphs in the document
    for para in doc.paragraphs:
        for run in para.runs:
            #replace_in_run(run)
            print(run.text)
    
    # Iterate through tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        #replace_in_run(run)
                        print(run.text)
    
    # Save the changes to the same file (in-place replacement)
    doc.save(doc_path)
    print(f"Document '{doc_path}' has been updated.")

# Example usage
#replace_text_in_place('./files/pdf.docx')
replace_text_in_docx()
